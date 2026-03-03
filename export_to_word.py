"""
Convert FINAL_REPORT.md to Word document format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def parse_markdown(md_file):
    """Parse markdown file and extract content"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def add_heading(doc, text, level):
    """Add heading with appropriate level"""
    # Remove markdown heading markers
    text = re.sub(r'^#+\s+', '', text)
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    """Add paragraph with optional style"""
    if style:
        para = doc.add_paragraph(text, style=style)
    else:
        para = doc.add_paragraph(text)
    return para

def convert_markdown_to_word(md_file, output_file):
    """Convert markdown to Word document"""
    print(f"Converting {md_file} to {output_file}...")
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown
    content = parse_markdown(md_file)
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_lines:
                    # Add code as formatted text
                    para = doc.add_paragraph()
                    run = para.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    para.paragraph_format.left_indent = Inches(0.5)
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.strip().startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and not line.strip().startswith('|'):
            # End of table - process it
            in_table = False
            if len(table_lines) >= 2:
                # Parse table
                rows = []
                for tline in table_lines:
                    if '---' in tline:  # Skip separator line
                        continue
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Create table in Word
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
                            # Bold header row
                            if row_idx == 0:
                                table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.bold = True
                    
                    doc.add_paragraph()  # Add spacing after table
            table_lines = []
        
        # Skip empty lines in specific contexts
        if not line.strip():
            if i > 0 and not lines[i-1].strip():
                i += 1
                continue
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            level = min(level, 3)  # Word supports up to 9, but limit to 3 for clarity
            heading_text = line.lstrip('#').strip()
            add_heading(doc, heading_text, level)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 60)
            i += 1
            continue
        
        # Handle bold/italic in text
        if '**' in line or '*' in line or '![' in line:
            # Handle images - just add as text reference
            if line.strip().startswith('!['):
                match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
                if match:
                    alt_text = match.group(1)
                    image_path = match.group(2)
                    para = doc.add_paragraph()
                    run = para.add_run(f"[Figure: {alt_text}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    para.add_run(f"\nPath: {image_path}")
                    i += 1
                    continue
            
            # Handle inline formatting
            para = doc.add_paragraph()
            # Simple bold/italic handling
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = para.add_run(part[1:-1])
                    run.font.italic = True
                else:
                    para.add_run(part)
            i += 1
            continue
        
        # Regular paragraph
        add_paragraph(doc, line)
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Document saved to: {output_file}")
    return output_file

if __name__ == "__main__":
    md_file = "FINAL_REPORT.md"
    output_file = "FINAL_REPORT.docx"
    
    if os.path.exists(md_file):
        convert_markdown_to_word(md_file, output_file)
        print(f"\nWord document created successfully!")
        print(f"Location: {os.path.abspath(output_file)}")
    else:
        print(f"Error: {md_file} not found!")
